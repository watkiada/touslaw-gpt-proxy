"""
Letter generation and saving service for AI-generated letters
"""
import os
import logging
import json
import asyncio
from typing import List, Dict, Any, Optional, Union
from datetime import datetime
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from pathlib import Path
import shutil

from app.core.config import settings
from app.services.ai.chat_service import AIChatService

logger = logging.getLogger(__name__)

class LetterService:
    """
    Service for generating and saving letters
    """
    
    def __init__(self):
        """
        Initialize letter service
        """
        self.chat_service = AIChatService()
        
        # Base directory for document storage
        self.base_dir = settings.DOCUMENT_STORAGE_PATH or "C:/WatkibotLawAssistant/documents"
    
    async def generate_letter(
        self,
        letter_type: str,
        recipient: str,
        subject: str,
        content_instructions: str,
        letterhead: Optional[Dict[str, str]] = None,
        case_id: Optional[int] = None,
        office_id: Optional[int] = None
    ) -> str:
        """
        Generate letter with AI
        
        Args:
            letter_type: Type of letter
            recipient: Letter recipient
            subject: Letter subject
            content_instructions: Instructions for letter content
            letterhead: Letterhead information (optional)
            case_id: Case ID (optional)
            office_id: Office ID (optional)
            
        Returns:
            Generated letter content
        """
        # Generate letter content using AI
        letter_content = await self.chat_service.generate_letter(
            letter_type=letter_type,
            recipient=recipient,
            subject=subject,
            content_instructions=content_instructions,
            letterhead=letterhead,
            case_id=case_id,
            office_id=office_id
        )
        
        return letter_content
    
    async def save_letter_to_case(
        self,
        letter_content: str,
        case_id: int,
        office_id: int,
        filename: str,
        format: str = "docx",
        letterhead: Optional[Dict[str, str]] = None
    ) -> Dict[str, Any]:
        """
        Save letter to case folder
        
        Args:
            letter_content: Letter content
            case_id: Case ID
            office_id: Office ID
            filename: Filename (without extension)
            format: File format (docx or pdf)
            letterhead: Letterhead information (optional)
            
        Returns:
            Dictionary with file information
        """
        # Validate format
        if format.lower() not in ["docx", "pdf"]:
            raise ValueError("Invalid format. Must be 'docx' or 'pdf'")
        
        # Ensure filename doesn't have extension
        filename = os.path.splitext(filename)[0]
        
        # Create case folder path
        case_folder = os.path.join(self.base_dir, f"office_{office_id}", f"case_{case_id}")
        os.makedirs(case_folder, exist_ok=True)
        
        # Generate file path
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(case_folder, f"{filename}_{timestamp}.{format.lower()}")
        
        # Create document based on format
        if format.lower() == "docx":
            # Create DOCX document
            doc_path = await self._create_docx_letter(
                letter_content=letter_content,
                output_path=file_path,
                letterhead=letterhead
            )
            
            return {
                "success": True,
                "file_path": doc_path,
                "format": "docx",
                "case_id": case_id,
                "office_id": office_id
            }
        else:
            # Create PDF document
            pdf_path = await self._create_pdf_letter(
                letter_content=letter_content,
                output_path=file_path,
                letterhead=letterhead
            )
            
            return {
                "success": True,
                "file_path": pdf_path,
                "format": "pdf",
                "case_id": case_id,
                "office_id": office_id
            }
    
    async def _create_docx_letter(
        self,
        letter_content: str,
        output_path: str,
        letterhead: Optional[Dict[str, str]] = None
    ) -> str:
        """
        Create DOCX letter
        
        Args:
            letter_content: Letter content
            output_path: Output file path
            letterhead: Letterhead information (optional)
            
        Returns:
            Path to created DOCX file
        """
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add letterhead if provided
        if letterhead:
            # Add firm name
            if "firm_name" in letterhead:
                heading = doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                firm_name = heading.add_run(letterhead["firm_name"])
                firm_name.bold = True
                firm_name.font.size = Pt(14)
            
            # Add address and contact info
            if any(key in letterhead for key in ["address", "city", "state", "zipCode", "phone", "email", "website"]):
                address_para = doc.add_paragraph()
                address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                address_parts = []
                if "address" in letterhead:
                    address_parts.append(letterhead["address"])
                
                location_parts = []
                if "city" in letterhead:
                    location_parts.append(letterhead["city"])
                if "state" in letterhead and "zipCode" in letterhead:
                    location_parts.append(f"{letterhead['state']} {letterhead['zipCode']}")
                elif "state" in letterhead:
                    location_parts.append(letterhead["state"])
                elif "zipCode" in letterhead:
                    location_parts.append(letterhead["zipCode"])
                
                if location_parts:
                    address_parts.append(", ".join(location_parts))
                
                if address_parts:
                    address_run = address_para.add_run("\n".join(address_parts))
                    address_run.font.size = Pt(10)
                
                contact_parts = []
                if "phone" in letterhead:
                    contact_parts.append(f"Phone: {letterhead['phone']}")
                if "email" in letterhead:
                    contact_parts.append(f"Email: {letterhead['email']}")
                if "website" in letterhead:
                    contact_parts.append(f"Website: {letterhead['website']}")
                
                if contact_parts:
                    if address_parts:
                        address_para.add_run("\n")
                    contact_run = address_para.add_run(" | ".join(contact_parts))
                    contact_run.font.size = Pt(10)
            
            # Add separator
            doc.add_paragraph("_" * 60)
        
        # Parse letter content and add to document
        lines = letter_content.split("\n")
        current_para = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line, add paragraph break
                current_para = None
            elif not current_para:
                # Start new paragraph
                current_para = doc.add_paragraph(line)
            else:
                # Continue current paragraph
                current_para.add_run("\n" + line)
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    async def _create_pdf_letter(
        self,
        letter_content: str,
        output_path: str,
        letterhead: Optional[Dict[str, str]] = None
    ) -> str:
        """
        Create PDF letter
        
        Args:
            letter_content: Letter content
            output_path: Output file path
            letterhead: Letterhead information (optional)
            
        Returns:
            Path to created PDF file
        """
        # First create DOCX, then convert to PDF
        # This approach ensures consistent formatting
        
        # Create temporary DOCX file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_docx_path = temp_file.name
        
        try:
            # Create DOCX
            await self._create_docx_letter(
                letter_content=letter_content,
                output_path=temp_docx_path,
                letterhead=letterhead
            )
            
            # Convert DOCX to PDF using LibreOffice (must be installed)
            # For Windows Server, we'll use a Python library that can convert without dependencies
            from docx2pdf import convert
            
            # Convert to PDF
            convert(temp_docx_path, output_path)
            
            return output_path
        
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_docx_path)
            except:
                pass
